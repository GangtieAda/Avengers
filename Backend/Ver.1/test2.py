from flask import Flask, request, jsonify, render_template
import mysql.connector
import os
import re
from werkzeug.utils import secure_filename
from docx import Document
import pdfplumber
import spacy
from dateutil import parser
from typing import List, Dict

app = Flask(__name__)

# Loading the SpaCy model
nlp = spacy.load("en_core_web_sm")

# Configuring MySQL Database Connection
db = mysql.connector.connect(
    host="localhost",
    user="root",
    password="Dqianru1211!",
    database="ResumeDatabase"
)
cursor = db.cursor()

# Upload folder and allowed file types
UPLOAD_FOLDER = 'uploads'
ALLOWED_EXTENSIONS = {'pdf', 'docx'}
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER

if not os.path.exists(UPLOAD_FOLDER):
    os.makedirs(UPLOAD_FOLDER)

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS



def extract_information_with_nlp(content: str) -> Dict[str, str]:
    """
    General resume information extraction function, adaptable to various layouts.
    """
    # Initializing fields
    result = {
        "first_name": "Unknown",
        "last_name": "Unknown",
        "email": "Unknown",
        "university": "Unknown",
        "major": "Unknown",
        "expected_graduation": "Unknown",
        "hiring_status": "pending",
    }

    # Block parsing
    sections = split_into_sections(content)

    # Extract Name
    result.update(extract_name(content))

    # Extract mailbox
    email_match = re.search(r"[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}", content)
    result["email"] = email_match.group() if email_match else "Unknown"

    # Extracting educational background
    education_section = sections.get("education", "")
    education_info = extract_education(education_section)
    result.update(education_info)

    return result


def split_into_sections(content: str) -> Dict[str, str]:
    """
    Divide your resume into sections based on headings.
    """
    # Common titles
    section_headers = [
        "education", "work experience", "skills", "projects", 
        "certifications", "languages", "awards", "hobbies"
    ]

    sections = {}
    current_section = "other"
    lines = content.splitlines()
    buffer = []

    for line in lines:
        line = line.strip()
        if not line:
            continue

        # Check if it is a title
        if any(header in line.lower() for header in section_headers):
            if buffer:
                sections[current_section] = "\n".join(buffer)
            current_section = line.lower()
            buffer = []
        else:
            buffer.append(line)

    if buffer:
        sections[current_section] = "\n".join(buffer)

    return sections


def extract_name(content: str) -> Dict[str, str]:
    """
    Extract names from content.
    """
    # Extracting Names Using SpaCy
    doc = nlp(content)
    name = None
    for ent in doc.ents:
        if ent.label_ == "PERSON":
            name = ent.text.strip()
            break

    # Split first and last name
    if name:
        names = name.split()
        first_name = names[0].capitalize()
        last_name = " ".join(names[1:]).capitalize() if len(names) > 1 else "Unknown"
    else:
        first_name = last_name = "Unknown"

    return {
        "first_name": first_name,
        "last_name": last_name,
    }


def extract_education(education_section: str) -> Dict[str, str]:
    """
    Extract educational background, including degree, major, and school information.
    """
    result = {
        "university": "Unknown",
        "major": "Unknown",
        "expected_graduation": "Unknown",
    }

    # Extract degree and major
    education_pattern = re.findall(
        r"(Bachelor|Master|PhD|Associate|Doctorate|B\.Sc\.|M\.Sc\.|M\.A\.|B\.A\.)\s*(?:of)?\s*([A-Za-z\s]+)?", 
        education_section
    )

    if education_pattern:
        degree, field = education_pattern[0]
        result["major"] = f"{degree} of {field}".strip()

    # Extract school name
    school_match = re.search(r"[A-Za-z\s]+University|Institute|College", education_section)
    if school_match:
        result["university"] = school_match.group().strip()

    # Extract expected graduation date
    grad_match = re.search(r"(\d{4})", education_section)
    if grad_match:
        result["expected_graduation"] = f"{grad_match.group()}-12-31"

    return result

def extract_text_from_pdf(file_path):
    """Extract PDF text using pdfplumber"""
    text = ""
    try:
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:  # Check if the page content is empty
                    text += page_text + "\n"
    except Exception as e:
        text = f"Error extracting text from PDF: {str(e)}"
    return text

def extract_text_from_docx(file_path):
    """Extract text from DOCX files"""
    text = ""
    try:
        doc = Document(file_path)
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text += paragraph.text + "\n"

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        text += cell_text + "\n"
    except Exception as e:
        text = f"Error extracting text from DOCX: {str(e)}"
    return text

@app.route('/')
def upload_page():
    return render_template('upload.html')

@app.route('/upload_resume', methods=['POST'])
def upload_resume():
    try:
        if 'resume' not in request.files:
            return jsonify({"error": "No file part"}), 400

        file = request.files['resume']
        document_type = request.form.get('document_type', 'Resume')
        ats_score = float(request.form.get('ats_score', 0))

        if file.filename == '':
            return jsonify({"error": "No selected file"}), 400

        if file and allowed_file(file.filename):
            filename = secure_filename(file.filename)
            file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            file.save(file_path)

            # Extract text content
            if filename.endswith('.pdf'):
                text_content = extract_text_from_pdf(file_path)
            elif filename.endswith('.docx'):
                text_content = extract_text_from_docx(file_path)
            else:
                return jsonify({"error": "Unsupported file format"}), 400

            # Extract key information
            info = extract_information_with_nlp(text_content)

            # Database insertion logic
            
            cursor.execute("""
                INSERT INTO Students (email, first_name, last_name, university, major, expected_graduation, hiring_status, created_at, updated_at)
                VALUES (%s, %s, %s, %s, %s, %s, %s, NOW(), NOW())
                ON DUPLICATE KEY UPDATE
                first_name = VALUES(first_name),
                last_name = VALUES(last_name),
                university = VALUES(university),
                major = VALUES(major),
                expected_graduation = VALUES(expected_graduation),
                hiring_status = VALUES(hiring_status),
                updated_at = NOW()
            """, (
                info['email'], 
                info['first_name'], 
                info['last_name'], 
                info['university'], 
                info['major'], 
                info['expected_graduation'], 
                info['hiring_status']
            ))
            db.commit()
            student_id = cursor.lastrowid

            cursor.execute("""
                INSERT INTO StudentDocuments (student_id, document_type, file_path, ats_score, upload_date)
                VALUES (%s, %s, %s, %s, NOW())
            """, (student_id, document_type, file_path, ats_score))
            db.commit()

            return jsonify({
                "message": "Resume uploaded and processed successfully!",
                "student_id": student_id,
                "data": info
            }), 201
        else:
            return jsonify({"error": "Invalid file type"}), 400
    except mysql.connector.Error as db_error:
        return jsonify({"error": f"Database error: {str(db_error)}"}), 500
    except Exception as e:
        return jsonify({"error": f"Unexpected error: {str(e)}"}), 500

if __name__ == '__main__':
    app.run(debug=True, port=5006)

# Pending questions: 1. resume template layout